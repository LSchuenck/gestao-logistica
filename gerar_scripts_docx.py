# -*- coding: utf-8 -*-
"""
Gera um .docx com todos os scripts do projeto Gestão Logística,
separados por camada: Config/SQL, Model, DAO, Controller, Views, Raiz.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r"c:\xampp\htdocs\Gestao Logistica"

# ── Estrutura das seções: (título, lista de caminhos relativos ordenados) ─────
SECTIONS = [
    ("Configuração e SQL", [
        r"config\conexao.php",
        r"config\auth.php",
        r"config\gestao_logistica.sql",
        r"config\migration_endereco.sql",
    ]),
    ("Models (Entidades)", [
        r"model\Transportadora.php",
        r"model\Motorista.php",
        r"model\Veiculo.php",
        r"model\Cliente.php",
        r"model\Armazem.php",
        r"model\Produto.php",
        r"model\Estoque.php",
        r"model\Rota.php",
        r"model\Viagem.php",
        r"model\Entrega.php",
        r"model\Frete.php",
        r"model\Usuario.php",
    ]),
    ("DAOs (Acesso a Dados)", [
        r"dao\EnderecoDao.php",
        r"dao\TransportadoraDao.php",
        r"dao\MotoristaDao.php",
        r"dao\VeiculoDao.php",
        r"dao\ClienteDao.php",
        r"dao\ArmazemDao.php",
        r"dao\ProdutoDao.php",
        r"dao\EstoqueDao.php",
        r"dao\MovimentacaoDao.php",
        r"dao\RotaDao.php",
        r"dao\ViagemDao.php",
        r"dao\EntregaDao.php",
        r"dao\FreteDao.php",
        r"dao\OperacaoDao.php",
        r"dao\AlertaDao.php",
        r"dao\IndicadorDao.php",
        r"dao\DashboardDao.php",
        r"dao\UsuarioDao.php",
    ]),
    ("Controllers (Lógica de Negócio)", [
        r"controller\AuthController.php",
        r"controller\TransportadoraController.php",
        r"controller\MotoristaController.php",
        r"controller\VeiculoController.php",
        r"controller\ClienteController.php",
        r"controller\ArmazemController.php",
        r"controller\ProdutoController.php",
        r"controller\EstoqueController.php",
        r"controller\RotaController.php",
        r"controller\ViagemController.php",
        r"controller\EntregaController.php",
        r"controller\FreteController.php",
        r"controller\OperacaoController.php",
        r"controller\AlertaController.php",
        r"controller\IndicadorController.php",
        r"controller\DashboardController.php",
        r"controller\UsuarioController.php",
    ]),
    ("Views (Interface)", [
        r"views\login.php",
        r"views\trocar_senha.php",
        r"views\index.php",
        r"views\transportadoras.php",
        r"views\motoristas.php",
        r"views\veiculos.php",
        r"views\clientes.php",
        r"views\armazens.php",
        r"views\produtos.php",
        r"views\estoque.php",
        r"views\rotas.php",
        r"views\viagens.php",
        r"views\entregas.php",
        r"views\frete.php",
        r"views\operacoes.php",
        r"views\alertas.php",
        r"views\indicadores.php",
        r"views\usuarios.php",
    ]),
    ("Páginas de Entrada (Raiz)", [
        r"index.php",
        r"login.php",
        r"logout.php",
        r"transportadoras.php",
        r"motoristas.php",
        r"veiculos.php",
        r"clientes.php",
        r"armazens.php",
        r"produtos.php",
        r"estoque.php",
        r"rotas.php",
        r"viagens.php",
        r"entregas.php",
        r"frete.php",
        r"operacoes.php",
        r"alertas.php",
        r"indicadores.php",
        r"usuarios.php",
        r"trocar_senha.php",
        r"instalar.php",
        r"gerar_hash.php",
    ]),
]

# ── Documento ─────────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def add_heading_section(text):
    """Título de seção (capítulo de camada)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)  # azul
    return p

def add_heading_file(rel_path):
    """Título do arquivo dentro da seção."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(rel_path.replace("\\", "/"))
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    # linha separadora via borda inferior
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BBBBBB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_code(content):
    """Adiciona o conteúdo do arquivo em fonte monospace."""
    # Divide em linhas para evitar parágrafos gigantes (melhor desempenho no Word)
    lines = content.splitlines()
    # Agrupa em blocos de 40 linhas para não criar milhares de parágrafos
    CHUNK = 40
    for i in range(0, len(lines), CHUNK):
        chunk_lines = lines[i:i+CHUNK]
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        run = p.add_run("\n".join(chunk_lines))
        run.font.name = "Courier New"
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    doc.add_paragraph()  # espaço após o bloco de código

# ── Capa ──────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(12)
r = p.add_run("Sistema de Gestão Logística")
r.bold = True
r.font.size = Pt(20)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(6)
r2 = p2.add_run("Listagem Completa de Scripts por Camada")
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(4)
r3 = p3.add_run("Lucas Schuenck · 2026")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Sumário simples
doc.add_page_break()
pt = doc.add_paragraph()
pt.paragraph_format.space_before = Pt(0)
pt.paragraph_format.space_after = Pt(10)
rt = pt.add_run("Índice de Seções")
rt.bold = True
rt.font.size = Pt(13)

for sec_title, files in SECTIONS:
    ps = doc.add_paragraph()
    ps.paragraph_format.space_after = Pt(2)
    rs = ps.add_run(f"  {sec_title}  ({len(files)} arquivo{'s' if len(files)!=1 else ''})")
    rs.font.size = Pt(11)
    for rel in files:
        full = os.path.join(BASE, rel)
        exists = os.path.isfile(full)
        pf = doc.add_paragraph()
        pf.paragraph_format.space_after = Pt(0)
        pf.paragraph_format.left_indent = Cm(1)
        icon = "✓" if exists else "✗"
        rf = pf.add_run(f"{icon}  {rel.replace(chr(92), '/')}")
        rf.font.size = Pt(9)
        rf.font.color.rgb = RGBColor(0x33, 0x88, 0x33) if exists else RGBColor(0xCC, 0x22, 0x22)

# ── Conteúdo ──────────────────────────────────────────────────────────────────
for sec_title, files in SECTIONS:
    doc.add_page_break()
    add_heading_section(f"▌ {sec_title}")

    for rel in files:
        full = os.path.join(BASE, rel)
        if not os.path.isfile(full):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run(f"[arquivo não encontrado: {rel}]")
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xCC, 0x22, 0x22)
            continue

        add_heading_file(rel)

        try:
            with open(full, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(full, "r", encoding="latin-1") as f:
                content = f.read()

        line_count = content.count("\n") + 1
        pm = doc.add_paragraph()
        pm.paragraph_format.space_before = Pt(0)
        pm.paragraph_format.space_after  = Pt(2)
        rm = pm.add_run(f"{line_count} linhas")
        rm.font.size = Pt(8)
        rm.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        rm.italic = True

        add_code(content)

# ── Salvar ────────────────────────────────────────────────────────────────────
output = os.path.join(BASE, "Scripts_Gestao_Logistica.docx")
doc.save(output)
print(f"Concluído! Arquivo salvo em:\n  {output}")
