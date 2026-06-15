# -*- coding: utf-8 -*-
"""
Gera a documentação de requisitos do Sistema de Gestão Logística no formato .docx,
seguindo o modelo acadêmico do IFSUDESTEMG - Campus Muriaé.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Margens ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def para(text="", bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         color=None, space_before=0, space_after=6, first_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if first_indent is not None:
        fmt.first_line_indent = Cm(first_indent)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def heading(text, level=1, size=12, bold=True, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=12, space_after=6):
    p = para(text, bold=bold, italic=italic, size=size, align=align,
             space_before=space_before, space_after=space_after)
    return p

def bullet(text, indent_cm=1.0, size=12):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Cabeçalho
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(11)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Linhas
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            cells[ci].paragraphs[0].runs[0].font.size = Pt(11)
    # Larguras
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CAPA
# ══════════════════════════════════════════════════════════════════════════════
para("")
para("")
para("INSTITUTO FEDERAL DE EDUCAÇÃO CIÊNCIA E TECNOLOGIA DO\nSUDESTE DE MINAS GERAIS - CAMPUS MURIAÉ\nGESTÃO DA TECNOLOGIA DA INFORMAÇÃO",
     bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=60)

para("Lucas Schuenck", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=120)

p = para("", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)
run1 = p.add_run("SISTEMA DE GESTÃO LOGÍSTICA:\n")
run1.bold = True
run1.font.size = Pt(12)
run2 = p.add_run("Documentação de Requisitos")
run2.bold = True
run2.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

para("", space_before=120, space_after=0)
para("Muriaé – MG\n2026", size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_before=0, space_after=0)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SUMÁRIO
# ══════════════════════════════════════════════════════════════════════════════
para("SUMÁRIO", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_before=0, space_after=12)

sumario = [
    ("1", "INTRODUÇÃO", "1"),
    ("1.1", "Público-Alvo", "1"),
    ("2", "DESCRIÇÃO GERAL DO PRODUTO", "2"),
    ("2.1", "Descrição do Cliente", "2"),
    ("2.2", "Descrição Geral do Produto", "3"),
    ("2.2.1", "Situação Atual da Empresa", "4"),
    ("2.2.2", "Escopo do Produto – Módulos e Funcionalidades", "5"),
    ("2.3", "Premissas", "8"),
    ("3", "REQUISITOS", "10"),
    ("3.1", "Requisitos Funcionais", "10"),
    ("3.2", "Requisitos Não Funcionais", "13"),
    ("3.3", "Regras de Negócio", "14"),
    ("4", "PLANO DE LIBERAÇÕES", "15"),
]
for num, titulo, pag in sumario:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{num}  {titulo}")
    r1.font.size = Pt(12)
    r1.bold = num in ("1", "2", "3", "4")
    # pontos e página
    tab = p.add_run(f" {'.' * max(1, 60 - len(num) - len(titulo))} {pag}")
    tab.font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 1 – INTRODUÇÃO
# ══════════════════════════════════════════════════════════════════════════════
heading("1  INTRODUÇÃO", size=12, bold=True, space_before=0)

para(
    "Este documento de requisitos tem como objetivo definir e descrever as funcionalidades "
    "e características do sistema denominado Sistema de Gestão Logística. Desenvolvido em "
    "PHP com banco de dados MySQL e rodando sobre infraestrutura XAMPP, o sistema visa "
    "centralizar e digitalizar as operações logísticas de empresas de transporte, abrangendo "
    "o controle de transportadoras, motoristas, veículos, clientes, armazéns, produtos, "
    "estoque, rotas, viagens, entregas, frete e alertas operacionais.",
    first_indent=1.25
)

para(
    "Este documento servirá de guia para os desenvolvedores e para todos os stakeholders "
    "envolvidos no projeto, garantindo o alinhamento entre os requisitos do cliente e a "
    "solução técnica a ser entregue.",
    first_indent=1.25
)

# 1.1
heading("1.1  Público-Alvo", size=12, bold=True, space_before=10)

para(
    "O público-alvo deste documento inclui:",
    first_indent=1.25
)
bullet("Equipe de Desenvolvimento: Desenvolvedores, programadores e demais membros responsáveis pela implementação do sistema.")
bullet("Gestores de Projeto: Gerentes e líderes de equipe que supervisionam o desenvolvimento e garantem o cumprimento dos requisitos.")
bullet("Clientes e Usuários Finais: Gestores logísticos, operadores de transportadora e demais usuários que interagirão diretamente com o sistema.")
bullet("Stakeholders: Professores orientadores, avaliadores acadêmicos e demais partes interessadas no projeto.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 2 – DESCRIÇÃO GERAL DO PRODUTO
# ══════════════════════════════════════════════════════════════════════════════
heading("2  DESCRIÇÃO GERAL DO PRODUTO", size=12, bold=True, space_before=0)

# 2.1
heading("2.1  Descrição do Cliente", size=12, bold=True, space_before=10)

para(
    "Durante o processo de levantamento de requisitos, identificou-se que empresas do setor "
    "logístico ainda realizam grande parte do controle operacional de forma manual, por meio de "
    "planilhas, anotações físicas e comunicação informal entre setores. Essa abordagem acarreta "
    "erros de comunicação, retrabalho, dificuldade no rastreamento de entregas e ausência de "
    "indicadores confiáveis para a tomada de decisão.",
    first_indent=1.25
)

para(
    "O cliente manifestou interesse em uma plataforma digital integrada que permita:",
    first_indent=1.25
)
bullet("Cadastrar e gerenciar transportadoras, motoristas, veículos e clientes.")
bullet("Controlar armazéns, produtos e movimentações de estoque.")
bullet("Planejar e acompanhar rotas, viagens e entregas em tempo real.")
bullet("Calcular fretes automaticamente com base nas rotas e veículos.")
bullet("Receber alertas automáticos sobre ocorrências críticas (atrasos, estoque baixo, CNH vencida).")
bullet("Visualizar indicadores de desempenho (KPIs) consolidados em um painel gerencial.")

# 2.2
heading("2.2  Descrição Geral do Produto", size=12, bold=True, space_before=10)

para(
    "O Sistema de Gestão Logística é uma aplicação web desenvolvida em PHP puro (sem "
    "framework), com banco de dados MySQL, interface responsiva em Bootstrap 5 e acesso "
    "restrito por autenticação de usuários com controle de perfis (ADMIN e GERENTE). O "
    "sistema é executado localmente via XAMPP e estruturado no padrão MVC (Model–View–"
    "Controller), com separação clara entre camadas de persistência (DAO), regras de negócio "
    "(Controller) e apresentação (View).",
    first_indent=1.25
)

para(
    "Seu objetivo principal é proporcionar visibilidade e controle completos sobre todas as "
    "etapas da cadeia logística: desde o cadastro de ativos (veículos, motoristas, armazéns) "
    "até o acompanhamento do ciclo de vida de uma entrega, passando pelo cálculo do frete e "
    "pela geração de alertas e indicadores.",
    first_indent=1.25
)

# 2.2.1
heading("2.2.1  Situação Atual da Empresa", size=12, bold=True, space_before=10)

para("A empresa enfrenta os seguintes desafios em suas operações logísticas:", first_indent=1.25)

heading("Controle Manual de Operações", size=12, bold=True, italic=True, space_before=8)
bullet("Registros de rotas e entregas realizados em planilhas ou papel.")
bullet("Comunicação informal entre motoristas e gestores, sujeita a erros.")
bullet("Impossibilidade de consultar o histórico de viagens de forma centralizada.")

heading("Rastreamento Limitado de Entregas", size=12, bold=True, italic=True, space_before=8)
bullet("Ausência de plataforma centralizada para acompanhamento do status das entregas.")
bullet("Clientes sem acesso a informações sobre suas encomendas.")
bullet("Atrasos identificados apenas após reclamação.")

heading("Gestão de Estoque Ineficiente", size=12, bold=True, italic=True, space_before=8)
bullet("Controle de produtos e armazéns feito manualmente.")
bullet("Risco de rupturas de estoque ou excesso de itens sem giro.")
bullet("Ausência de alertas automáticos para reposição.")

heading("Ausência de Indicadores", size=12, bold=True, italic=True, space_before=8)
bullet("Falta de métricas consolidadas sobre desempenho de entregas, frota e motoristas.")
bullet("Tomada de decisão prejudicada pela ausência de dados em tempo real.")
bullet("Dificuldade em identificar gargalos operacionais.")

# 2.2.2
heading("2.2.2  Escopo do Produto – Módulos e Funcionalidades", size=12, bold=True, space_before=10)

para(
    "O sistema abrange os seguintes módulos e funcionalidades principais:",
    first_indent=1.25
)

heading("Módulo de Cadastros Operacionais", size=12, bold=True, italic=True, space_before=8)
add_table(
    ["Funcionalidade", "Descrição"],
    [
        ["Transportadoras", "Cadastro de empresas transportadoras com CNPJ, endereço, contato e status (Ativa/Inativa)."],
        ["Motoristas", "Cadastro de motoristas vinculados a transportadoras, com CPF, CNH, categoria e validade."],
        ["Veículos", "Cadastro de veículos com placa, modelo, tipo, capacidade e status (Disponível/Em Viagem/Manutenção)."],
        ["Clientes", "Cadastro de clientes com CPF/CNPJ, telefone e endereço de entrega."],
        ["Armazéns", "Cadastro de centros de distribuição com nome e endereço completo."],
        ["Produtos", "Cadastro de produtos com nome, unidade e quantidade mínima de estoque."],
    ],
    col_widths=[5, 11]
)

heading("Módulo de Estoque", size=12, bold=True, italic=True, space_before=8)
add_table(
    ["Funcionalidade", "Descrição"],
    [
        ["Controle de Estoque", "Consulta de saldo por produto e armazém com localização."],
        ["Movimentações", "Registro de entradas e saídas de produtos com rastreamento de tipo (entrada/saída/transferência)."],
        ["Alertas de Estoque", "Geração automática de alertas quando o saldo cai abaixo do mínimo cadastrado."],
    ],
    col_widths=[5, 11]
)

heading("Módulo de Operações Logísticas", size=12, bold=True, italic=True, space_before=8)
add_table(
    ["Funcionalidade", "Descrição"],
    [
        ["Rotas", "Planejamento de rotas com origem, destino, distância e tempo estimado."],
        ["Viagens", "Registro de viagens vinculando motorista, veículo e rota, com status e datas."],
        ["Entregas", "Controle de entregas vinculadas a clientes, com status e rastreamento."],
        ["Frete", "Cálculo e registro de valores de frete por viagem, com tipo de cobrança."],
    ],
    col_widths=[5, 11]
)

heading("Módulo Gerencial", size=12, bold=True, italic=True, space_before=8)
add_table(
    ["Funcionalidade", "Descrição"],
    [
        ["Alertas", "Painel de alertas operacionais em 4 tipos: ATRASO (entregas vencidas), VIAGEM (prazo de chegada ultrapassado), ESTOQUE (armazém com menos de 10 unidades) e DESVIO_ROTA (desvios registrados em viagens)."],
        ["Indicadores (KPIs)", "Dashboard com métricas de entregas, frota, motoristas e estoque."],
        ["Usuários", "Gestão de usuários do sistema com perfis de acesso (ADMIN/GERENTE)."],
    ],
    col_widths=[5, 11]
)

# 2.3
heading("2.3  Premissas", size=12, bold=True, space_before=10)

para("As premissas a seguir foram identificadas durante o levantamento de requisitos:", first_indent=1.25)

heading("Infraestrutura Local", size=12, bold=True, italic=True, space_before=8)
para("O sistema será executado em servidor local com XAMPP (Apache + MySQL). Acesso via navegador web na rede interna.", first_indent=1.25)

heading("Equipe de Administração", size=12, bold=True, italic=True, space_before=8)
para("Haverá ao menos um usuário com perfil ADMIN responsável pelo cadastro inicial e manutenção dos dados.", first_indent=1.25)

heading("Conectividade para ViaCEP", size=12, bold=True, italic=True, space_before=8)
para("A busca automática de endereço por CEP requer conexão com a internet para consulta à API pública ViaCEP.", first_indent=1.25)

heading("Integração com Power Automate para Envio de E-mails", size=12, bold=True, italic=True, space_before=8)
para("O envio automático de e-mails com senha temporária depende de um fluxo configurado no Microsoft Power Automate (Microsoft Flow), acionado via webhook HTTP. Caso o servidor não possua conexão com a internet ou o fluxo esteja inativo, a senha temporária será exibida na tela do administrador como fallback.", first_indent=1.25)

heading("Treinamento dos Usuários", size=12, bold=True, italic=True, space_before=8)
para("Os usuários serão capacitados antes da implantação para garantir uso correto das funcionalidades.", first_indent=1.25)

heading("Backup dos Dados", size=12, bold=True, italic=True, space_before=8)
para("Deverá ser realizado backup regular do banco de dados MySQL para prevenção de perdas.", first_indent=1.25)

heading("Navegadores Compatíveis", size=12, bold=True, italic=True, space_before=8)
para("O sistema é compatível com navegadores modernos (Chrome, Firefox, Edge) em suas versões recentes.", first_indent=1.25)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 3 – REQUISITOS
# ══════════════════════════════════════════════════════════════════════════════
heading("3  REQUISITOS", size=12, bold=True, space_before=0)

# 3.1
heading("3.1  Requisitos Funcionais", size=12, bold=True, space_before=10)

rfs = [
    ("RF01 – Autenticação de Usuários", [
        "O sistema deve exigir login com e-mail e senha para acesso.",
        "As senhas devem ser armazenadas com hash seguro (bcrypt ou equivalente).",
        "Ao criar um novo usuário, o sistema deve gerar uma senha temporária automaticamente e enviá-la ao e-mail do usuário via API do Power Automate (Microsoft Flow). Caso o envio falhe, a senha temporária deve ser exibida na tela como fallback.",
        "O usuário deve ser obrigado a trocar a senha no primeiro acesso.",
        "Deve existir função de troca de senha pelo próprio usuário autenticado.",
        "O sistema deve suportar três perfis: ADMIN (acesso total), GERENTE (acesso operacional) e OPERADOR.",
    ]),
    ("RF02 – Gestão de Transportadoras", [
        "O sistema deve permitir cadastro, visualização, edição e exclusão de transportadoras.",
        "Cada transportadora deve possuir: CNPJ, razão social, nome fantasia, telefone, e-mail, endereço e status (ATIVA/INATIVA).",
        "Transportadoras com motoristas ou veículos vinculados não podem ser excluídas.",
        "O status pode ser alternado entre ATIVA e INATIVA.",
    ]),
    ("RF03 – Gestão de Motoristas", [
        "O sistema deve permitir cadastro, visualização, edição e exclusão de motoristas.",
        "Cada motorista deve possuir: transportadora vinculada, nome, CPF, CNH, categoria (B/C/D/E), validade da CNH, telefone e status (ATIVO/INATIVO).",
        "O sistema deve alertar automaticamente quando a CNH estiver vencida.",
        "CPF e número de CNH devem ser únicos no sistema.",
    ]),
    ("RF04 – Gestão de Veículos", [
        "O sistema deve permitir cadastro, visualização, edição e exclusão de veículos.",
        "Cada veículo deve possuir: transportadora vinculada, placa, modelo, tipo (Van/Caminhão/Carreta/Bitrem), capacidade de carga (kg) e status (DISPONIVEL/EM_VIAGEM/MANUTENCAO).",
        "A placa deve ser única e armazenada em letras maiúsculas.",
        "O status pode ser atualizado diretamente na listagem via dropdown.",
    ]),
    ("RF05 – Gestão de Clientes", [
        "O sistema deve permitir cadastro, visualização, edição e exclusão de clientes.",
        "Cada cliente deve possuir: nome/razão social, CPF/CNPJ (opcional), telefone e endereço de entrega.",
        "Clientes com entregas vinculadas não podem ser excluídos.",
        "O sistema deve exibir o total de entregas por cliente na listagem.",
    ]),
    ("RF06 – Gestão de Armazéns", [
        "O sistema deve permitir cadastro, visualização, edição e exclusão de armazéns/centros de distribuição.",
        "Cada armazém deve possuir: nome e endereço completo (CEP, logradouro, número, bairro, cidade, estado).",
        "O sistema deve exibir o total de SKUs e quantidade de itens em estoque por armazém.",
        "Armazéns com produtos localizados não podem ser excluídos.",
    ]),
    ("RF07 – Gestão de Produtos e Estoque", [
        "O sistema deve permitir cadastro, visualização, edição e exclusão de produtos.",
        "O controle de estoque deve indicar quantidade disponível por produto em cada armazém.",
        "Movimentações de estoque (entrada/saída) devem ser registradas com data e responsável.",
        "O sistema deve gerar alerta automático quando o saldo de um produto em qualquer armazém específico cair abaixo de 10 unidades.",
        "O estoque é descontado automaticamente ao iniciar uma viagem e estornado automaticamente caso a viagem seja cancelada.",
    ]),
    ("RF08 – Gestão de Rotas", [
        "O sistema deve permitir cadastro, visualização, edição e exclusão de rotas.",
        "Cada rota deve possuir: nome, origem, destino, distância (km) e tempo estimado.",
        "As rotas devem poder ser associadas a múltiplas viagens.",
    ]),
    ("RF09 – Gestão de Viagens", [
        "O sistema deve permitir cadastro, visualização, edição e exclusão de viagens.",
        "Cada viagem deve vincular: rota, motorista, veículo, data de início, data prevista de término e status.",
        "Uma viagem pode conter múltiplas entregas associadas.",
        "O sistema deve registrar o histórico de rastreamento (posição/evento) por viagem.",
    ]),
    ("RF10 – Gestão de Entregas", [
        "O sistema deve permitir cadastro, visualização e atualização do status das entregas.",
        "Cada entrega deve vincular: cliente, viagem e produtos entregues.",
        "Os status possíveis são: PENDENTE, EM_TRANSITO, ENTREGUE, CANCELADA.",
        "O sistema deve exibir o histórico de entregas por cliente.",
    ]),
    ("RF11 – Cálculo e Gestão de Frete", [
        "O sistema deve permitir o registro do valor de frete por viagem.",
        "O frete deve conter: viagem vinculada, tipo de cobrança (por km, por peso ou fixo), valor base e valor total calculado.",
        "A transportadora do frete deve ser determinada automaticamente a partir da viagem selecionada, não sendo possível alterá-la manualmente.",
        "O sistema deve exibir o total de faturamento de frete por período.",
    ]),
    ("RF12 – Painel de Alertas", [
        "O sistema deve exibir alertas ativos agrupados por tipo e criticidade.",
        "Os alertas são classificados em quatro tipos: ATRASO (entregas com data prevista vencida e não entregues), VIAGEM (viagens em trânsito com prazo de chegada ultrapassado), ESTOQUE (armazéns com menos de 10 unidades de um produto) e DESVIO_ROTA (desvios registrados manualmente durante viagens).",
        "Os alertas de ATRASO, VIAGEM e ESTOQUE são gerados dinamicamente a cada acesso à tela; os de DESVIO_ROTA são persistidos no banco via módulo de Operações.",
        "O badge de alertas no menu lateral exibe o total de alertas ativos de todos os tipos, atualizando a contagem a cada carregamento de página.",
    ]),
    ("RF13 – Painel de Indicadores (KPIs)", [
        "O sistema deve exibir um dashboard com os principais indicadores logísticos.",
        "Os KPIs devem incluir: entregas pendentes, em trânsito e atrasadas; viagens ativas; total de alertas ativos (todos os tipos); veículos disponíveis; motoristas ativos; e valor total de fretes do mês.",
        "O card de Alertas Ativos deve refletir o somatório de todos os alertas ativos (ATRASO + VIAGEM + ESTOQUE + DESVIO_ROTA), não apenas registros persistidos.",
        "Os indicadores devem ser atualizados em tempo real a partir dos dados do banco.",
    ]),
]

for titulo, itens in rfs:
    heading(titulo, size=12, bold=True, space_before=8)
    for item in itens:
        bullet(item)

# 3.2
heading("3.2  Requisitos Não Funcionais", size=12, bold=True, space_before=12)

rnfs = [
    ("RNF01 – Segurança", [
        "Todas as senhas devem ser armazenadas com hash (password_hash do PHP com BCRYPT).",
        "O acesso a qualquer página do sistema deve exigir autenticação ativa via sessão PHP.",
        "Perfis de acesso devem restringir funcionalidades conforme o nível do usuário (ADMIN/GERENTE).",
        "Dados de entrada devem ser sanitizados para prevenir SQL Injection e XSS.",
    ]),
    ("RNF02 – Desempenho", [
        "As páginas de listagem devem carregar em menos de 2 segundos em rede local.",
        "Consultas ao banco de dados devem utilizar prepared statements via PDO.",
        "Operações que envolvem múltiplas tabelas devem ser executadas em transações atômicas.",
    ]),
    ("RNF03 – Usabilidade", [
        "A interface deve ser responsiva e funcionar em telas de desktop e tablet.",
        "Formulários devem aplicar máscaras automáticas (CNPJ, CPF, telefone, CEP).",
        "O CEP deve disparar busca automática de endereço via API ViaCEP.",
        "Mensagens de erro e confirmação devem ser claras e objetivas.",
    ]),
    ("RNF04 – Disponibilidade", [
        "O sistema deve estar disponível durante o horário comercial (07h–22h) sem interrupções não programadas.",
        "Deve existir um procedimento de backup diário automático do banco de dados MySQL.",
    ]),
    ("RNF05 – Manutenibilidade", [
        "O código deve seguir o padrão MVC com separação entre Model, View e Controller.",
        "DAOs devem centralizar toda a lógica de acesso ao banco de dados.",
        "O banco de dados deve ter um script SQL autoritativo e atualizado (gestao_logistica.sql).",
    ]),
    ("RNF06 – Compatibilidade", [
        "O sistema deve funcionar nos navegadores Google Chrome, Mozilla Firefox e Microsoft Edge (versões dos últimos 2 anos).",
        "O ambiente de execução é PHP 8.x + MySQL 8.x + Apache (XAMPP).",
    ]),
    ("RNF07 – Escalabilidade", [
        "A arquitetura MVC deve permitir a adição de novos módulos sem impacto nos existentes.",
        "O banco de dados deve suportar crescimento incremental de registros sem necessidade de reestruturação.",
    ]),
]

for titulo, itens in rnfs:
    heading(titulo, size=12, bold=True, space_before=8)
    for item in itens:
        bullet(item)

# 3.3
heading("3.3  Regras de Negócio", size=12, bold=True, space_before=12)

add_table(
    ["ID", "Regra de Negócio"],
    [
        ["RN01", "Transportadoras com motoristas ou veículos vinculados não podem ser excluídas."],
        ["RN02", "Motoristas com rotas ou viagens vinculadas não podem ser excluídos."],
        ["RN03", "Veículos com rotas ou viagens vinculadas não podem ser excluídos."],
        ["RN04", "Clientes com entregas vinculadas não podem ser excluídos."],
        ["RN05", "Armazéns com produtos em estoque não podem ser excluídos."],
        ["RN06", "O número de CNH e o CPF do motorista devem ser únicos no sistema."],
        ["RN07", "A placa do veículo deve ser única e armazenada em letras maiúsculas."],
        ["RN08", "Apenas motoristas ATIVOS e veículos DISPONÍVEIS podem ser vinculados a novas viagens."],
        ["RN09", "Alertas de CNH devem ser gerados automaticamente para motoristas com validade vencida ou a vencer em 30 dias."],
        ["RN10", "Alertas de estoque são gerados quando o saldo de um produto em um armazém específico cair abaixo de 10 unidades. A verificação é feita por armazém individualmente, não pelo somatório total do produto."],
        ["RN11", "O valor total do frete deve ser calculado conforme o tipo de cobrança: por km (distância × valor/km), por peso (peso × valor/kg) ou fixo."],
        ["RN12", "Uma entrega só pode ter status ENTREGUE se a viagem vinculada estiver com status CONCLUÍDA."],
        ["RN13", "Somente usuários com perfil ADMIN podem criar, editar ou excluir outros usuários do sistema."],
        ["RN14", "A busca de endereço por CEP depende de conexão com a internet para consulta à API ViaCEP."],
        ["RN15", "O estoque dos produtos de uma entrega é descontado automaticamente ao iniciar a viagem (transição para EM_TRANSITO). Caso a viagem seja cancelada, o estoque é estornado automaticamente."],
        ["RN16", "A transportadora de um frete é determinada automaticamente pela viagem selecionada (viagem → rota → motorista → transportadora), não podendo ser alterada manualmente."],
        ["RN17", "Ao criar um novo usuário, o sistema gera uma senha temporária aleatória, armazena o hash bcrypt e dispara um e-mail via webhook do Power Automate. O usuário é obrigado a trocar a senha no primeiro acesso."],
    ],
    col_widths=[2.5, 13.5]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 4 – PLANO DE LIBERAÇÕES
# ══════════════════════════════════════════════════════════════════════════════
heading("4  PLANO DE LIBERAÇÕES", size=12, bold=True, space_before=0)

para(
    "O plano de liberações descreve as iterações de desenvolvimento do sistema, definindo "
    "quais funcionalidades serão entregues em cada fase, com seus respectivos objetivos e "
    "atividades principais.",
    first_indent=1.25
)

heading("Release 1 – Fundação do Sistema", size=12, bold=True, space_before=10)

iteracoes = [
    (
        "Iteração 1: Configuração e Infraestrutura",
        "01/03/2026", "15/03/2026",
        "Estabelecer a estrutura do projeto e configurar o ambiente de desenvolvimento.",
        [
            "Configuração do ambiente XAMPP (PHP 8, MySQL 8, Apache).",
            "Criação do banco de dados com script SQL (gestao_logistica.sql).",
            "Implementação do sistema de autenticação com sessões PHP e hash de senhas.",
            "Criação da estrutura MVC: pastas model/, dao/, controller/, views/.",
            "Implementação da barra de navegação e layout base responsivo (Bootstrap 5).",
        ]
    ),
    (
        "Iteração 2: Módulos de Cadastro – Transportadoras, Motoristas e Veículos",
        "16/03/2026", "31/03/2026",
        "Implementar os cadastros da cadeia de transporte, que são a base referencial do sistema.",
        [
            "CRUD completo de Transportadoras com endereço e controle de status.",
            "CRUD completo de Motoristas com validação de CNH e vínculo com transportadora.",
            "CRUD completo de Veículos com controle de status e capacidade de carga.",
            "Validações de integridade referencial (impedimento de exclusão com vínculos).",
            "Busca automática de CEP via ViaCEP nos formulários de endereço.",
        ]
    ),
    (
        "Iteração 3: Módulos de Cadastro – Clientes, Armazéns e Produtos",
        "01/04/2026", "15/04/2026",
        "Implementar os cadastros de clientes, infraestrutura de armazenagem e produtos.",
        [
            "CRUD completo de Clientes com endereço de entrega.",
            "CRUD completo de Armazéns com métricas de estoque (SKUs e total de itens).",
            "CRUD completo de Produtos com quantidade mínima de estoque.",
            "Módulo de Estoque: localização de produtos por armazém e movimentações.",
        ]
    ),
]

for titulo, data_ini, data_fim, objetivo, atividades in iteracoes:
    heading(titulo, size=12, bold=True, italic=True, space_before=8)
    p = doc.add_paragraph()
    r = p.add_run(f"Data Inicial: {data_ini}    Data Final: {data_fim}")
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Objetivo: ")
    r2.bold = True
    r2.font.size = Pt(12)
    r3 = p2.add_run(objetivo)
    r3.font.size = Pt(12)
    p2.paragraph_format.space_after = Pt(4)
    p3 = doc.add_paragraph()
    p3.add_run("Atividades:").bold = True
    p3.runs[0].font.size = Pt(12)
    p3.paragraph_format.space_after = Pt(4)
    for at in atividades:
        bullet(at)

heading("Release 2 – Operações Logísticas", size=12, bold=True, space_before=10)

iteracoes2 = [
    (
        "Iteração 4: Rotas, Viagens e Rastreamento",
        "16/04/2026", "30/04/2026",
        "Implementar os módulos de planejamento e execução de operações logísticas.",
        [
            "CRUD completo de Rotas com origem, destino e distância.",
            "CRUD completo de Viagens com vínculo a rota, motorista e veículo.",
            "Controle de status de viagem (Planejada / Em Andamento / Concluída / Cancelada).",
            "Registro de rastreamento por viagem (histórico de eventos e posições).",
        ]
    ),
    (
        "Iteração 5: Entregas e Frete",
        "01/05/2026", "15/05/2026",
        "Implementar o ciclo completo de entregas e o cálculo de frete.",
        [
            "Módulo de Entregas: vínculo com viagem e cliente, controle de status e produtos entregues.",
            "Módulo de Frete: cálculo por km, por peso ou valor fixo, e registro por viagem.",
            "Painel de entregas por status (Pendente / Em Trânsito / Entregue / Cancelada).",
            "Histórico de entregas por cliente.",
        ]
    ),
    (
        "Iteração 6: Alertas, KPIs e Refinamentos",
        "16/05/2026", "31/05/2026",
        "Implementar o painel gerencial com alertas e indicadores de desempenho.",
        [
            "Módulo de Alertas: geração automática de alertas para CNH vencida, estoque baixo e entregas atrasadas.",
            "Dashboard de Indicadores (KPIs): métricas de entregas, frota, estoque e faturamento.",
            "Funcionalidade de edição de registros em todos os módulos (Transportadoras, Motoristas, Veículos, Clientes, Armazéns).",
            "Refinamentos de usabilidade e correção de bugs reportados.",
            "Testes de integração e validação final do sistema.",
        ]
    ),
]

for titulo, data_ini, data_fim, objetivo, atividades in iteracoes2:
    heading(titulo, size=12, bold=True, italic=True, space_before=8)
    p = doc.add_paragraph()
    r = p.add_run(f"Data Inicial: {data_ini}    Data Final: {data_fim}")
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Objetivo: ")
    r2.bold = True
    r2.font.size = Pt(12)
    r3 = p2.add_run(objetivo)
    r3.font.size = Pt(12)
    p2.paragraph_format.space_after = Pt(4)
    p3 = doc.add_paragraph()
    p3.add_run("Atividades:").bold = True
    p3.runs[0].font.size = Pt(12)
    p3.paragraph_format.space_after = Pt(4)
    for at in atividades:
        bullet(at)

# ── Salvar ────────────────────────────────────────────────────────────────────
output = r"c:\xampp\htdocs\Gestao Logistica\Documentacao_Gestao_Logistica.docx"
doc.save(output)
print(f"Documento gerado com sucesso: {output}")
